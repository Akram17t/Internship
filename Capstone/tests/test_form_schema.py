from __future__ import annotations

import io
import json
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import fitz
from fastapi.testclient import TestClient

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from backend.api.forms_service import (  # noqa: E402
    extract_docx_placeholder_fields,
    fill_docx_template_placeholders,
    fill_schema_form,
)
from backend.api.main import app  # noqa: E402


TARGET_FORM_PATH = "Form - Perjalanan Dinas - Penyelesaian (Template).pdf"
EXIT_INTERVIEW_FORM_PATH = "Form - Exit Interview (Template).pdf"
OTHER_FORM_PATH = "Form - Missing Schema (Template).pdf"


def _sample_png_bytes() -> bytes:
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.Rect(0, 0, 20, 10), False)
    pixmap.clear_with(0xFFFFFF)
    return pixmap.tobytes("png")


PNG_BYTES = _sample_png_bytes()


def _normalized_pdf_text(text: str) -> str:
    return text.replace("\u00a0", " ").replace("\u2010", "-")


class FormSchemaRouteTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.client = TestClient(app)

    def _valid_values(self) -> dict[str, str]:
        return {
            "divisi": "Finance",
            "nama": "Akram",
            "jabatan": "Staff",
            "tujuan_kota_daerah": "Bandung",
            "lama": "2 hari, 10-11 Juli 2026",
            "total_biaya": "1500000",
        }

    def test_schema_endpoint_returns_target_schema(self) -> None:
        response = self.client.get("/api/forms/schema", params={"path": TARGET_FORM_PATH})

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["path"], TARGET_FORM_PATH)
        self.assertEqual(payload["pages"][0]["width"], 612)
        self.assertTrue(any(field["id"] == "tanda_tangan_pemohon" for field in payload["fields"]))
        division_field = next(field for field in payload["fields"] if field["id"] == "divisi")
        self.assertTrue(division_field["clear"])
        self.assertGreaterEqual(division_field["clear_padding"], 0)

    def test_all_schema_fields_are_preview_positioned(self) -> None:
        schema_dir = PROJECT_ROOT / "backend" / "form_schemas"
        for schema_path in schema_dir.glob("*.json"):
            with self.subTest(schema=schema_path.name):
                payload = json.loads(schema_path.read_text(encoding="utf-8"))
                response = self.client.get("/api/forms/schema", params={"path": payload["path"]})
                self.assertEqual(response.status_code, 200)
                fields = response.json()["fields"]
                self.assertTrue(fields)
                self.assertFalse(
                    [
                        field["label"]
                        for field in fields
                        if field.get("page") is None or not field.get("rect")
                    ]
                )

    def test_frontend_has_no_extra_field_section(self) -> None:
        frontend_files = [
            PROJECT_ROOT / "frontend" / "web" / "assets" / "js" / "forms.js",
            PROJECT_ROOT / "frontend" / "web" / "assets" / "styles.css",
            PROJECT_ROOT / "frontend" / "web" / "index.html",
        ]
        combined = "\n".join(path.read_text(encoding="utf-8") for path in frontend_files)
        self.assertNotIn("Field tambahan", combined)
        self.assertNotIn("form-extra-fields", combined)
        self.assertNotIn("renderExtraSchemaFields", combined)

    def test_schema_endpoint_404_for_form_without_schema(self) -> None:
        response = self.client.get("/api/forms/schema", params={"path": OTHER_FORM_PATH})

        self.assertEqual(response.status_code, 404)

    def test_schema_endpoint_returns_exit_interview_schema(self) -> None:
        response = self.client.get(
            "/api/forms/schema",
            params={"path": EXIT_INTERVIEW_FORM_PATH},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["path"], EXIT_INTERVIEW_FORM_PATH)
        self.assertEqual(len(payload["pages"]), 2)
        self.assertTrue(any(field["id"] == "kompensasi_benefit" for field in payload["fields"]))
        self.assertTrue(any((field.get("layout") or {}).get("kind") == "choice_matrix" for field in payload["fields"]))
        self.assertTrue(any(field["type"] == "signature_image" for field in payload["fields"]))

    def test_fill_endpoint_rejects_unknown_schema_field(self) -> None:
        response = self.client.post(
            "/api/forms/fill",
            json={
                "path": TARGET_FORM_PATH,
                "values": {**self._valid_values(), "bogus_field": "x"},
            },
        )

        self.assertEqual(response.status_code, 422)
        self.assertIn("tidak dikenal", response.json()["detail"])

    def test_fill_endpoint_allows_missing_form_values(self) -> None:
        response = self.client.post(
            "/api/forms/fill",
            json={
                "path": TARGET_FORM_PATH,
                "values": {"divisi": "Finance"},
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_fill_endpoint_rejects_non_image_signature_upload(self) -> None:
        payload = {"path": TARGET_FORM_PATH, "values": self._valid_values()}
        response = self.client.post(
            "/api/forms/fill",
            data={"payload": json.dumps(payload)},
            files={"tanda_tangan_pemohon": ("signature.txt", io.BytesIO(b"not-image"), "text/plain")},
        )

        self.assertEqual(response.status_code, 422)
        self.assertIn("tanda tangan belum didukung", response.json()["detail"])

    def test_fill_endpoint_rejects_pdf_for_schema_form(self) -> None:
        values = {
            **self._valid_values(),
            "hotel": "2 malam @ Rp 500.000 x 1 kamar",
            "baris_2_keterangan": "1000000",
            "nama_2": "Akram",
            "jabatan_2": "Staff",
        }
        payload = {"path": TARGET_FORM_PATH, "values": values, "output_format": "pdf"}

        response = self.client.post(
            "/api/forms/fill",
            json=payload,
        )

        self.assertEqual(response.status_code, 422)
        self.assertIn("PDF terisi belum didukung", response.json()["detail"])

    def test_fill_endpoint_returns_docx_for_schema_form(self) -> None:
        try:
            from docx import Document
            import pdf2docx  # noqa: F401
        except ImportError:
            self.skipTest("DOCX converter dependencies are not installed in this interpreter.")

        response = self.client.post(
            "/api/forms/fill",
            json={
                "path": TARGET_FORM_PATH,
                "values": {"divisi": "Finance", "nama": "Akram", "jabatan": "Staff"},
                "output_format": "docx",
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        document = Document(io.BytesIO(response.content))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.extend(paragraph.text for paragraph in cell.paragraphs)
        text = "\n".join(chunks)
        self.assertIn("Finance", text)
        self.assertIn("Akram", text)

    def test_fill_endpoint_returns_docx_for_exit_interview_schema(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed in this interpreter.")

        response = self.client.post(
            "/api/forms/fill",
            json={
                "path": EXIT_INTERVIEW_FORM_PATH,
                "values": {
                    "no_form": "EI-001",
                    "nama_karyawan": "Akram",
                    "nik": "EMP-001",
                    "jabatan_terakhir": "Software Engineer",
                    "departemen": "Professional Service",
                    "atasan_langsung": "Budi",
                    "tanggal_mengajukan_resign": "2026-07-10",
                    "tanggal_terakhir_kerja": "2026-07-31",
                    "nama_perusahaan_baru": "Contoso",
                    "jenis_industri": "Teknologi",
                    "posisi": "Senior Engineer",
                    "tanggal_bergabung": "2026-08-15",
                    "kompensasi_benefit": True,
                    "lain_lain": True,
                    "textarea_37_alasan_penjelasan": "Mencari tantangan baru dan kompensasi yang lebih baik.",
                    "textarea_38_4_apakah_ada_hal_yang_dapat_dilakukan_perusahaan_untuk_mempertahankan_anda": "Skema kerja hybrid mungkin bisa membantu.",
                    "saya_memahami_peran_dan_tanggung_jawab_pekerjaan_saya_setuju": True,
                    "beban_kerja_sesuai_dengan_kapasitas_dan_waktu_kerja_tidak_setuju": True,
                    "atasan_memberikan_arahan_dan_dukungan_yang_memadai_setuju": True,
                    "fasilitas_kerja_mendukung_produktivitas_sangat_setuju": True,
                    "kompensasi_dan_benefit_sesuai_tanggung_jawab_pekerjaan_tidak_setuju": True,
                    "kesempatan_pengembangan_karir_tersedia_dengan_baik_tidak_setuju": True,
                    "komunikasi_internal_perusahaan_berjalan_dengan_baik_setuju": True,
                    "textarea_58_6_hal_yang_perlu_dipertahankan_ditingkatkan_perusahaan": "Pertahankan budaya tim yang suportif.",
                    "nama": "Akram",
                    "jabatan": "HR",
                    "nama_2": "Budi",
                    "jabatan_2": "Manager",
                    "nama_3": "Citra",
                    "jabatan_3": "HRBP",
                },
            },
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.headers["content-type"],
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        document = Document(io.BytesIO(response.content))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.extend(paragraph.text for paragraph in cell.paragraphs)
        text = _normalized_pdf_text("\n".join(chunks))
        self.assertIn("EI-001", text)
        self.assertIn("Professional Service", text)
        self.assertIn("Contoso", text)
        self.assertIn("Akram", text)
        self.assertIn("X", text)
        self.assertIn("Pertahankan budaya tim yang suportif.", text)
        self.assertIn("Manager", text)


class SchemaRenderUnitTests(unittest.TestCase):
    def test_extract_docx_placeholders_infers_inline_table_and_signature_labels(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "template.docx"
            document = Document()
            document.add_paragraph("Nama: [ ] Jabatan: [ ]")

            table = document.add_table(rows=2, cols=4)
            table.cell(0, 0).text = "No"
            table.cell(0, 1).text = "Item"
            table.cell(0, 2).text = "Ya"
            table.cell(0, 3).text = "Tidak"
            table.cell(1, 1).text = "Gaji Terakhir"
            table.cell(1, 2).text = "[ ]"
            table.cell(1, 3).text = "[ ]"

            signature = document.add_table(rows=2, cols=3)
            signature.cell(0, 0).text = "Karyawan"
            signature.cell(0, 1).text = "HRGA Dept. Head"
            signature.cell(0, 2).text = "Diketahui Atasan"
            signature.cell(1, 0).text = "Nama: [ ] Jabatan: [ ]"
            signature.cell(1, 1).text = "Nama: [ ] Jabatan: [ ]"
            signature.cell(1, 2).text = "Nama: [ ] Jabatan: [ ]"
            document.save(path)

            fields = extract_docx_placeholder_fields(path)

        labels = [field["label"] for field in fields]
        self.assertIn("Nama", labels)
        self.assertIn("Jabatan", labels)
        self.assertIn("Gaji Terakhir - Ya", labels)
        self.assertIn("Gaji Terakhir - Tidak", labels)
        self.assertIn("Karyawan - Nama", labels)
        self.assertIn("HRGA Dept. Head - Jabatan", labels)

    def test_fill_docx_template_replaces_placeholder_split_across_runs(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "template.docx"
            output_path = Path(temp_dir) / "filled.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Nama: ")
            paragraph.add_run("[")
            paragraph.add_run(" ]")
            paragraph.add_run(" Jabatan: ")
            paragraph.add_run("[ ]")
            document.save(path)

            content = fill_docx_template_placeholders(
                path,
                {"nama": "Akram", "jabatan": "Professional Service"},
            )
            output_path.write_bytes(content)
            filled = Document(output_path)

        text = "\n".join(paragraph.text for paragraph in filled.paragraphs)
        self.assertIn("Akram", text)
        self.assertIn("Professional Service", text)
        self.assertNotIn("[ ]", text)

    def test_fill_schema_form_supports_textarea_checkbox_and_signature(self) -> None:
        schema = {
            "path": "test.pdf",
            "title": "Synthetic",
            "pages": [{"number": 0, "width": 200.0, "height": 200.0}],
            "fields": [
                {
                    "id": "notes",
                    "label": "Catatan",
                    "type": "textarea",
                    "page": 0,
                    "rect": {"x": 20.0, "y": 20.0, "width": 120.0, "height": 40.0},
                    "required": False,
                    "section": "Test",
                    "placeholder": None,
                    "font_size": 10.0,
                    "align": "left",
                    "line_height": 1.08,
                    "clear": True,
                    "clear_padding": 1.0,
                },
                {
                    "id": "approved",
                    "label": "Setuju",
                    "type": "checkbox",
                    "page": 0,
                    "rect": {"x": 20.0, "y": 80.0, "width": 26.0, "height": 18.0},
                    "required": False,
                    "section": "Test",
                    "placeholder": None,
                    "font_size": 10.0,
                    "align": "center",
                    "line_height": 1.0,
                    "clear": True,
                    "clear_padding": 1.0,
                },
                {
                    "id": "signature",
                    "label": "Signature",
                    "type": "signature_image",
                    "page": 0,
                    "rect": {"x": 20.0, "y": 110.0, "width": 40.0, "height": 24.0},
                    "required": False,
                    "section": "Test",
                    "placeholder": None,
                    "font_size": 10.0,
                    "align": "left",
                    "line_height": 1.0,
                    "clear": True,
                    "clear_padding": 1.0,
                },
            ],
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "test.pdf"
            doc = fitz.open()
            doc.new_page(width=200, height=200)
            doc.save(path)
            doc.close()

            with patch("backend.api.forms_service._schema_for_resolved_path", return_value=schema):
                content = fill_schema_form(
                    path,
                    {"notes": "Baris pertama\nBaris kedua", "approved": True},
                    {
                        "signature": {
                            "filename": "signature.png",
                            "content_type": "image/png",
                            "content": PNG_BYTES,
                        }
                    },
                )

        with fitz.open(stream=content, filetype="pdf") as rendered:
            page = rendered[0]
            text = _normalized_pdf_text(page.get_text())
            self.assertIn("Baris pertama", text)
            self.assertIn("X", text)
            self.assertGreater(len(page.get_images(full=True)), 0)


if __name__ == "__main__":
    unittest.main()
