from __future__ import annotations

import json
import re
import tempfile
from copy import deepcopy
from hashlib import sha1
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from fastapi import HTTPException

from backend.api.storage import (
    _document_kind_for_path,
    _get_data_dir,
    _resolve_document_path,
)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Cocokkan placeholder bracket yang dipakai template form, misalnya "[  ]".
FORM_PLACEHOLDER_PATTERN = re.compile(r"\[[^\[\]]*\]")
# Field segment adalah segmen teks yang seluruh isinya hanya satu bracket.
# Placeholder inline seperti "Nama: [ ]" atau header dilewati agar form tetap singkat.
FORM_FIELD_CELL_PATTERN = re.compile(r"^\[[^\[\]]*\]$")

FORM_FIELD_TYPES = {"text", "textarea", "date", "checkbox", "signature_image"}
IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/jpg"}
WORD_FIELD_SECTION = "Word"
WORD_PLACEHOLDER_KIND = "docx_placeholder"
WORD_CELL_KIND = "docx_cell"
WORD_FILLABLE_HEADERS = {"ya", "tidak", "tanggal serah terima", "keterangan"}
_WORD_TEMPLATE_CACHE_DIR = Path(__file__).resolve().parents[1] / "cache" / "form_word_templates"

# Pakai Calibri agar hasil isi form mendekati template Office; fallback ke Helvetica
# kalau backend jalan di environment yang tidak punya font Windows.
_CALIBRI_FONT_PATH = Path("C:/Windows/Fonts/calibri.ttf")
_FILL_FONT = "Calibri" if _CALIBRI_FONT_PATH.exists() else "helv"
_FILL_FONT_FILE = str(_CALIBRI_FONT_PATH) if _CALIBRI_FONT_PATH.exists() else None
_FILL_FONT_OBJECT = fitz.Font(fontfile=_FILL_FONT_FILE) if _FILL_FONT_FILE else None
_MIN_FONT_SIZE = 6.0
_ALIGNMENTS = {"left": 0, "center": 1, "right": 2}


def _text_length(value: str, *, fontsize: float) -> float:
    if _FILL_FONT_OBJECT is not None:
        return _FILL_FONT_OBJECT.text_length(value, fontsize=fontsize)
    return fitz.get_text_length(value, fontname=_FILL_FONT, fontsize=fontsize)


def _convert_pdf_to_docx_file(pdf_path: Path, docx_path: Path) -> None:
    try:
        from pdf2docx import Converter
    except ImportError as error:
        raise HTTPException(
            status_code=500,
            detail="Converter Word belum terpasang. Jalankan pip install -r requirements.txt.",
        ) from error

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    converter = Converter(str(pdf_path))
    try:
        converter.convert(str(docx_path), start=0, end=None)
    finally:
        converter.close()

    if not docx_path.exists() or docx_path.stat().st_size <= 0:
        raise HTTPException(status_code=500, detail="Form Word gagal dibuat.")


def filled_pdf_to_docx(pdf_content: bytes) -> bytes:
    """Convert PDF terisi ke DOCX dengan pdf2docx agar bukan image-only document."""
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        pdf_path = temp_path / "filled-form.pdf"
        docx_path = temp_path / "filled-form.docx"
        pdf_path.write_bytes(pdf_content)
        _convert_pdf_to_docx_file(pdf_path, docx_path)
        return docx_path.read_bytes()

def _word_cache_path(pdf_path: Path) -> Path:
    stat = pdf_path.stat()
    try:
        relative = _relative_form_path(pdf_path)
    except ValueError:
        relative = pdf_path.resolve().as_posix()
    digest = sha1(f"{relative}:{stat.st_mtime_ns}:{stat.st_size}".encode("utf-8")).hexdigest()
    safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", pdf_path.stem).strip("._") or "form"
    return _WORD_TEMPLATE_CACHE_DIR / f"{safe_stem}-{digest}.docx"


def word_template_docx_path(pdf_path: Path) -> Path:
    """Return cached DOCX template converted from the source PDF."""
    if pdf_path.suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Template ini bukan PDF.")
    cache_path = _word_cache_path(pdf_path)
    if cache_path.exists() and cache_path.stat().st_size > 0:
        return cache_path
    _convert_pdf_to_docx_file(pdf_path, cache_path)
    return cache_path


def get_word_template_docx(pdf_path: Path) -> bytes:
    return word_template_docx_path(pdf_path).read_bytes()


def _page_segments(page, page_number: int) -> list[dict]:
    """Ambil semua segmen teks PDF beserta bbox dari satu halaman."""
    segments: list[dict] = []
    for block in page.get_text("rawdict")["blocks"]:
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            text = "".join(char["c"] for span in spans for char in span.get("chars", []))
            if not text.strip():
                continue
            segments.append(
                {
                    "page": page_number,
                    "text": text.strip(),
                    "bbox": tuple(line["bbox"]),
                    "size": spans[0]["size"] if spans else 10.0,
                }
            )
    return segments


def _clean_label(text: str) -> str | None:
    # Rapikan teks kandidat label; buang jika kosong atau masih mengandung bracket.
    candidate = text.strip().rstrip(":").strip()
    if candidate and not FORM_PLACEHOLDER_PATTERN.search(candidate):
        return candidate
    return None


def _clean_docx_label(text: str) -> str | None:
    candidate = FORM_PLACEHOLDER_PATTERN.sub(" ", text or "")
    candidate = re.sub(r"\s+", " ", candidate).strip(" \t\r\n:;-")
    if not candidate:
        return None
    return candidate


def _slugify_field_id(label: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", label.lower()).strip("_")
    return slug or "field"


def _field_label_key(label: str) -> str:
    label = re.sub(r"\([^)]*\)", " ", label)
    key = re.sub(r"[^a-z0-9]+", " ", label.lower()).strip()
    key = re.sub(r"\bno\b", "nomor", key)
    key = key.replace("tujuan perjalanan dinas", "tujuan kota daerah")
    key = re.sub(r"\s+", " ", key)
    return key


def _unique_field_id(base_id: str, used_ids: set[str]) -> str:
    candidate = base_id
    index = 2
    while candidate in used_ids:
        candidate = f"{base_id}_{index}"
        index += 1
    used_ids.add(candidate)
    return candidate


def _join_label_parts(*parts: str | None) -> str | None:
    cleaned: list[str] = []
    for part in parts:
        label = _clean_docx_label(part or "")
        if label and label not in cleaned:
            cleaned.append(label)
    if not cleaned:
        return None
    return " - ".join(cleaned)


def _inline_docx_label(text: str, start: int) -> str | None:
    before = text[:start]
    segment = FORM_PLACEHOLDER_PATTERN.split(before)[-1]
    return _clean_docx_label(segment)


def _docx_cell_text(cell: Any) -> str:
    return " ".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _docx_table_context(table: Any, row_index: int, column_index: int) -> dict[str, str | None]:
    rows = table.rows
    row_cells = rows[row_index].cells
    row_label = None
    for index in range(column_index - 1, -1, -1):
        raw_label = _docx_cell_text(row_cells[index])
        if FORM_PLACEHOLDER_PATTERN.search(raw_label):
            continue
        row_label = _clean_docx_label(raw_label)
        if row_label:
            break

    column_parts: list[str] = []
    for index in range(row_index - 1, -1, -1):
        cells = rows[index].cells
        if column_index >= len(cells):
            continue
        raw_label = _docx_cell_text(cells[column_index])
        if FORM_PLACEHOLDER_PATTERN.search(raw_label):
            continue
        column_label_part = _clean_docx_label(raw_label)
        if column_label_part and column_label_part not in column_parts:
            column_parts.insert(0, column_label_part)

    return {"row_label": row_label, "column_label": " ".join(column_parts) or None}


def _load_word_document(docx_path: Path) -> Any:
    try:
        from docx import Document
    except ImportError as error:
        raise HTTPException(
            status_code=500,
            detail="Library python-docx belum terpasang. Jalankan pip install -r requirements.txt.",
        ) from error
    return Document(str(docx_path))


def _iter_docx_blocks(document: Any):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def _append_docx_fields_from_text(
    fields: list[dict[str, Any]],
    used_ids: set[str],
    text: str,
    *,
    context: dict[str, str | None] | None = None,
) -> None:
    if not text:
        return
    context = context or {}
    for match in FORM_PLACEHOLDER_PATTERN.finditer(text):
        placeholder = match.group(0)
        inside = placeholder[1:-1].strip()
        inline_label = _inline_docx_label(text, match.start())
        row_label = context.get("row_label")
        column_label = context.get("column_label")

        if inside:
            label = inside
        elif column_label and inline_label:
            label = _join_label_parts(column_label, inline_label)
        elif row_label and column_label:
            label = _join_label_parts(row_label, column_label)
        else:
            label = inline_label or _join_label_parts(row_label, column_label)
        if not label:
            label = f"Field {len(fields) + 1}"
        if re.fullmatch(r"\d+\s*-\s*.+", label):
            continue

        base_id = _slugify_field_id(label)
        field_id = _unique_field_id(base_id, used_ids)
        fields.append(
            {
                "id": field_id,
                "label": label,
                "type": "text",
                "page": None,
                "rect": None,
                "required": False,
                "section": WORD_FIELD_SECTION,
                "placeholder": label,
                "font_size": 10.0,
                "align": "left",
                "line_height": 1.08,
                "clear": False,
                "clear_padding": 0.0,
                "layout": {"kind": WORD_PLACEHOLDER_KIND},
                "docx_placeholder": placeholder,
            }
        )


def _is_repeated_section_row(cells: list[Any]) -> bool:
    labels = [_clean_docx_label(_docx_cell_text(cell)) for cell in cells]
    non_empty = [label for label in labels if label]
    return bool(non_empty) and len(set(non_empty)) == 1 and len(non_empty) > 1


def _docx_table_header_info(
    table: Any,
    previous: dict[str, Any] | None = None,
) -> dict[str, Any] | None:
    for row_index, row in enumerate(table.rows):
        labels = [_clean_docx_label(_docx_cell_text(cell)) or "" for cell in row.cells]
        lowered = [label.lower() for label in labels]
        if "item" not in lowered:
            continue
        fillable_columns = {
            index: labels[index]
            for index, label in enumerate(lowered)
            if label in WORD_FILLABLE_HEADERS
        }
        if not fillable_columns:
            continue
        return {
            "header_row": row_index,
            "item_column": lowered.index("item"),
            "columns": fillable_columns,
        }

    if previous and table.rows and len(table.rows[0].cells) >= len(previous.get("labels", [])):
        return {
            "header_row": -1,
            "item_column": previous["item_column"],
            "columns": previous["columns"],
        }
    return None


def _header_snapshot(table: Any, info: dict[str, Any]) -> dict[str, Any] | None:
    header_row = int(info.get("header_row", -1))
    if header_row < 0:
        return None
    labels = [_clean_docx_label(_docx_cell_text(cell)) or "" for cell in table.rows[header_row].cells]
    return {
        "labels": labels,
        "item_column": info["item_column"],
        "columns": info["columns"],
    }


def _signature_schema_fields(schema: dict[str, Any]) -> list[dict[str, Any]]:
    fields = []
    for field in schema["fields"]:
        layout = field.get("layout") if isinstance(field.get("layout"), dict) else {}
        rect = field.get("rect")
        if layout.get("kind") == "signature_block" and isinstance(rect, dict):
            fields.append(field)
    return sorted(fields, key=lambda field: (int(field.get("page") or 0), float(field["rect"]["x"])))


def _role_key(label: str) -> str:
    return _field_label_key(label).replace("tanda tangan ", "")


def _signature_text_field_for_docx_label(schema: dict[str, Any], label: str) -> dict[str, Any] | None:
    if " - " not in label:
        return None
    role_label, value_label = [part.strip() for part in label.rsplit(" - ", 1)]
    value_key = _field_label_key(value_label)
    if value_key not in {"nama", "jabatan"}:
        return None

    signatures = _signature_schema_fields(schema)
    signature = next(
        (
            candidate
            for candidate in signatures
            if _role_key(str((candidate.get("layout") or {}).get("row_label") or candidate.get("label") or ""))
            == _role_key(role_label)
        ),
        None,
    )
    if signature is None:
        return None

    signature_rect = signature["rect"]
    signature_center_x = float(signature_rect["x"]) + float(signature_rect["width"]) / 2
    candidates = []
    for field in schema["fields"]:
        layout = field.get("layout") if isinstance(field.get("layout"), dict) else {}
        rect = field.get("rect")
        if (
            str(layout.get("kind") or "") == "inline_placeholder"
            and isinstance(rect, dict)
            and int(field.get("page") or 0) == int(signature.get("page") or 0)
            and _field_label_key(str(field.get("label") or "")) == value_key
            and float(rect["y"]) >= float(signature_rect["y"])
        ):
            field_center_x = float(rect["x"]) + float(rect["width"]) / 2
            candidates.append((abs(field_center_x - signature_center_x), field))
    if not candidates:
        return None
    candidates.sort(key=lambda item: item[0])
    return candidates[0][1]


def _is_anonymous_docx_field(field: dict[str, Any]) -> bool:
    return bool(
        field.get("docx_placeholder")
        and re.fullmatch(r"Field\s+\d+", str(field.get("label") or ""), flags=re.IGNORECASE)
    )


def _infer_named_placeholder_pdf_rect(field: dict[str, Any], pdf_path: Path) -> None:
    if _field_has_preview_rect(field):
        return
    placeholder = str(field.get("docx_placeholder") or "")
    inside = placeholder.strip()[1:-1].strip() if placeholder.startswith("[") and placeholder.endswith("]") else ""
    if not inside:
        return

    doc = fitz.open(pdf_path)
    try:
        for page_number, page in enumerate(doc):
            rects = page.search_for(f"[{inside}]") or page.search_for(inside)
            if not rects:
                continue
            rect = sorted(rects, key=lambda item: (item.y0, item.x0))[0]
            field["page"] = page_number
            field["rect"] = {
                "x": float(rect.x0),
                "y": float(rect.y0),
                "width": float(rect.width),
                "height": float(rect.height),
            }
            field["clear"] = True
            field["clear_padding"] = 0.5
            return
    finally:
        doc.close()


def _infer_docx_table_cell_pdf_rect(field: dict[str, Any], schema: dict[str, Any], pdf_path: Path) -> None:
    if _field_has_preview_rect(field):
        return
    layout = field.get("layout") if isinstance(field.get("layout"), dict) else {}
    if layout.get("kind") != WORD_CELL_KIND:
        return
    row_label = str(layout.get("row_label") or "").strip()
    column_label = str(layout.get("column_label") or "").strip()
    if not row_label or not column_label:
        return

    doc = fitz.open(pdf_path)
    try:
        row_matches: list[tuple[int, fitz.Rect]] = []
        for page_number, page in enumerate(doc):
            for rect in page.search_for(row_label):
                row_matches.append((page_number, rect))
        if not row_matches:
            return
        page_number, row_rect = sorted(row_matches, key=lambda item: (item[0], item[1].y0, item[1].x0))[0]
    finally:
        doc.close()

    column_fields = []
    for schema_field in schema["fields"]:
        schema_layout = schema_field.get("layout") if isinstance(schema_field.get("layout"), dict) else {}
        rect = schema_field.get("rect")
        if (
            isinstance(rect, dict)
            and str(schema_layout.get("column_label") or "").strip().lower() == column_label.lower()
            and str(schema_layout.get("kind") or "") in {"table_cell", "choice_matrix"}
        ):
            column_fields.append(schema_field)
    if not column_fields:
        return

    same_page = [candidate for candidate in column_fields if int(candidate.get("page") or 0) == page_number]
    candidates = same_page or column_fields
    candidates.sort(key=lambda candidate: abs(float(candidate["rect"]["y"]) - float(row_rect.y0)))
    column_rect = candidates[0]["rect"]
    field["page"] = page_number
    field["rect"] = {
        "x": float(column_rect["x"]),
        "y": float(row_rect.y0),
        "width": float(column_rect["width"]),
        "height": float(column_rect["height"]),
    }
    field["font_size"] = float(candidates[0].get("font_size") or 10)
    field["align"] = str(candidates[0].get("align") or "left")
    field["line_height"] = float(candidates[0].get("line_height") or 1.08)
    field["clear"] = bool(candidates[0].get("clear", True))
    field["clear_padding"] = max(float(candidates[0].get("clear_padding") or 1.0), 0.0)


def _infer_docx_field_pdf_rect(field: dict[str, Any], schema: dict[str, Any], pdf_path: Path) -> None:
    if field.get("rect") or field.get("page") is not None:
        return

    _infer_named_placeholder_pdf_rect(field, pdf_path)
    if _field_has_preview_rect(field):
        return

    _infer_docx_table_cell_pdf_rect(field, schema, pdf_path)
    if _field_has_preview_rect(field):
        return

    label = str(field.get("label") or "")
    if " - " not in label:
        return
    role_label, value_label = [part.strip() for part in label.rsplit(" - ", 1)]
    value_key = _field_label_key(value_label)
    if value_key not in {"nama", "jabatan"}:
        return

    signature_fields = _signature_schema_fields(schema)
    if not signature_fields:
        return
    page_number = None
    signature = None
    for candidate in signature_fields:
        candidate_role = str((candidate.get("layout") or {}).get("row_label") or candidate.get("label") or "")
        if _role_key(candidate_role) == _role_key(role_label):
            signature = candidate
            page_number = int(candidate.get("page") or 0)
            break
    if signature is None or page_number is None:
        return

    page_signature_fields = [
        candidate for candidate in signature_fields if int(candidate.get("page") or 0) == page_number
    ]
    signature_index = page_signature_fields.index(signature)

    doc = fitz.open(pdf_path)
    try:
        page = doc[page_number]
        signature_rect = _rect_from_schema(signature["rect"])
        label_rects = sorted(
            [
                rect
                for rect in page.search_for(f"{value_label}:")
                if rect.y0 >= signature_rect.y0 + signature_rect.height * 0.55
            ],
            key=lambda rect: (rect.y0, rect.x0),
        )
        if not label_rects:
            return

        baseline_y = label_rects[0].y0
        row_rects = sorted(
            [rect for rect in label_rects if abs(rect.y0 - baseline_y) < 3],
            key=lambda rect: rect.x0,
        )
        if signature_index >= len(row_rects):
            return
        label_rect = row_rects[signature_index]

        next_role_x = (
            float(page_signature_fields[signature_index + 1]["rect"]["x"])
            if signature_index + 1 < len(page_signature_fields)
            else page.rect.x1 - 48
        )
        if value_key == "nama":
            jabatan_rects = sorted(
                [
                    rect
                    for rect in page.search_for("Jabatan:")
                    if abs(rect.y0 - label_rect.y0) < 3
                ],
                key=lambda rect: rect.x0,
            )
            right_edge = jabatan_rects[signature_index].x0 - 2 if signature_index < len(jabatan_rects) else next_role_x - 12
        else:
            right_edge = next_role_x - 12

        x = label_rect.x1 + 2
        field["page"] = page_number
        field["rect"] = {
            "x": float(x),
            "y": float(label_rect.y0 - 1),
            "width": float(max(28.0, right_edge - x)),
            "height": float(label_rect.height + 3),
        }
        field["font_size"] = float(signature.get("font_size") or 10)
        field["align"] = "left"
        field["clear"] = True
        field["clear_padding"] = 0.5
    finally:
        doc.close()


def _append_docx_blank_cell_fields(
    fields: list[dict[str, Any]],
    used_ids: set[str],
    table: Any,
    table_index: int,
    header_info: dict[str, Any] | None,
) -> None:
    if not header_info:
        return

    start_row = max(int(header_info.get("header_row", -1)) + 1, 0)
    item_column = int(header_info["item_column"])
    fillable_columns: dict[int, str] = header_info["columns"]
    for row_index, row in enumerate(table.rows[start_row:], start=start_row):
        cells = row.cells
        if item_column >= len(cells) or _is_repeated_section_row(list(cells)):
            continue
        item_label = _clean_docx_label(_docx_cell_text(cells[item_column]))
        if not item_label:
            continue
        if re.fullmatch(r"\d+", item_label):
            continue

        for column_index, column_label in fillable_columns.items():
            if column_index >= len(cells):
                continue
            if _clean_docx_label(_docx_cell_text(cells[column_index])):
                continue
            label = _join_label_parts(item_label, column_label)
            if not label:
                continue
            field_type = "checkbox" if column_label.lower() in {"ya", "tidak"} else "text"
            field_id = _unique_field_id(_slugify_field_id(label), used_ids)
            fields.append(
                {
                    "id": field_id,
                    "label": label,
                    "type": field_type,
                    "page": None,
                    "rect": None,
                    "required": False,
                    "section": WORD_FIELD_SECTION,
                    "placeholder": label,
                    "font_size": 10.0,
                    "align": "left",
                    "line_height": 1.08,
                    "clear": False,
                    "clear_padding": 0.0,
                    "layout": {
                        "kind": WORD_CELL_KIND,
                        "row_label": item_label,
                        "column_label": column_label,
                    },
                    "docx_target": {
                        "table": table_index,
                        "row": row_index,
                        "column": column_index,
                    },
                }
            )


def _append_docx_labelled_blank_fields(
    fields: list[dict[str, Any]],
    used_ids: set[str],
    table: Any,
    table_index: int,
) -> None:
    for row_index in range(len(table.rows) - 1):
        label_cells = list(table.rows[row_index].cells)
        target_cells = list(table.rows[row_index + 1].cells)
        if not target_cells:
            continue
        label = _clean_docx_label(_docx_cell_text(label_cells[0])) if label_cells else None
        if not label:
            continue
        if re.fullmatch(r"\d+(?:\s*-\s*.+)?", label):
            continue
        label_texts = [_clean_docx_label(_docx_cell_text(cell)) or "" for cell in label_cells]
        non_empty_labels = [text for text in label_texts if text]
        if not non_empty_labels or len(set(non_empty_labels)) > 1:
            continue
        if any(_clean_docx_label(_docx_cell_text(cell)) for cell in target_cells):
            continue
        if not (label.endswith(":") or "?" in label or len(label) >= 20):
            continue
        field_id = _unique_field_id(_slugify_field_id(label), used_ids)
        fields.append(
            {
                "id": field_id,
                "label": label.rstrip(":").strip(),
                "type": "textarea",
                "page": None,
                "rect": None,
                "required": False,
                "section": WORD_FIELD_SECTION,
                "placeholder": label.rstrip(":").strip(),
                "font_size": 10.0,
                "align": "left",
                "line_height": 1.08,
                "clear": False,
                "clear_padding": 0.0,
                "layout": {"kind": WORD_CELL_KIND},
                "docx_target": {
                    "table": table_index,
                    "row": row_index + 1,
                    "column": 0,
                },
            }
        )


def extract_docx_placeholder_fields(docx_path: Path, *, used_ids: set[str] | None = None) -> list[dict[str, Any]]:
    document = _load_word_document(docx_path)
    fields: list[dict[str, Any]] = []
    reserved_ids = set(used_ids or set())
    table_index = 0
    previous_header: dict[str, Any] | None = None

    for kind, block in _iter_docx_blocks(document):
        if kind == "paragraph":
            _append_docx_fields_from_text(fields, reserved_ids, block.text)
            continue

        header_info = _docx_table_header_info(block, previous_header)
        _append_docx_blank_cell_fields(fields, reserved_ids, block, table_index, header_info)
        _append_docx_labelled_blank_fields(fields, reserved_ids, block, table_index)
        snapshot = _header_snapshot(block, header_info) if header_info else None
        if snapshot:
            previous_header = snapshot

        seen_cells: set[int] = set()
        for row_index, row in enumerate(block.rows):
            for column_index, cell in enumerate(row.cells):
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                context = _docx_table_context(block, row_index, column_index)
                for paragraph in cell.paragraphs:
                    _append_docx_fields_from_text(
                        fields,
                        reserved_ids,
                        paragraph.text,
                        context=context,
                    )
        table_index += 1
    return fields


def _docx_fields_for_schema(schema: dict[str, Any], pdf_path: Path) -> list[dict[str, Any]]:
    docx_path = word_template_docx_path(pdf_path)
    docx_fields = extract_docx_placeholder_fields(docx_path)

    existing_by_label: dict[str, dict[str, Any]] = {}
    for field in schema["fields"]:
        label_key = _field_label_key(str(field.get("label") or ""))
        if label_key and label_key not in existing_by_label:
            existing_by_label[label_key] = field
    label_mapped_ids: set[str] = set()
    for raw_field in docx_fields:
        if _is_anonymous_docx_field(raw_field):
            continue
        signature_existing = _signature_text_field_for_docx_label(schema, str(raw_field.get("label") or ""))
        if signature_existing:
            label_mapped_ids.add(str(signature_existing["id"]))
            continue
        field_key = _field_label_key(str(raw_field.get("label") or ""))
        existing = existing_by_label.get(field_key)
        if existing is None:
            existing = next(
                (
                    candidate
                    for label_key, candidate in existing_by_label.items()
                    if label_key and field_key and (label_key in field_key or field_key in label_key)
                ),
                None,
            )
        if existing:
            label_mapped_ids.add(str(existing["id"]))

    anonymous_checkbox_queue = [
        field
        for field in sorted(
            schema["fields"],
            key=lambda item: (
                int(item.get("page") or 0),
                float((item.get("rect") or {}).get("y", 0)),
                float((item.get("rect") or {}).get("x", 0)),
            ),
        )
        if str(field.get("type")) == "checkbox" and str(field["id"]) not in label_mapped_ids
    ]
    coordinate_queues: dict[str, list[dict[str, Any]]] = {}
    for field in schema["fields"]:
        layout = field.get("layout") if isinstance(field.get("layout"), dict) else {}
        column_label = str(layout.get("column_label") or "").strip()
        if not column_label or column_label.lower() not in WORD_FILLABLE_HEADERS:
            continue
        coordinate_queues.setdefault(column_label.lower(), []).append(field)

    mapped_fields: list[dict[str, Any]] = []
    used_ids = {str(field["id"]) for field in schema["fields"]}
    for raw_field in docx_fields:
        field = deepcopy(raw_field)
        layout = field.get("layout") if isinstance(field.get("layout"), dict) else {}
        if layout.get("kind") == WORD_CELL_KIND and layout.get("column_label"):
            column_key = str(layout.get("column_label") or "").strip().lower()
            queue = coordinate_queues.get(column_key) or []
            if queue:
                field["id"] = str(queue.pop(0)["id"])
        elif _is_anonymous_docx_field(field) and anonymous_checkbox_queue:
            field["id"] = str(anonymous_checkbox_queue.pop(0)["id"])
        else:
            signature_existing = _signature_text_field_for_docx_label(schema, str(field.get("label") or ""))
            field_key = _field_label_key(str(field.get("label") or ""))
            existing = signature_existing or existing_by_label.get(field_key)
            if existing is None:
                existing = next(
                    (
                        candidate
                        for label_key, candidate in existing_by_label.items()
                        if label_key and field_key and (label_key in field_key or field_key in label_key)
                    ),
                    None,
                )
            if existing:
                field["id"] = str(existing["id"])

        if str(field["id"]) not in used_ids:
            field["id"] = _unique_field_id(str(field["id"]), used_ids)
        _infer_docx_field_pdf_rect(field, schema, pdf_path)
        mapped_fields.append(field)
    return mapped_fields


def _schema_with_docx_fields(schema: dict[str, Any], pdf_path: Path) -> dict[str, Any]:
    merged = deepcopy(schema)
    used_ids = {str(field["id"]) for field in merged["fields"]}
    existing_labels = {str(field.get("label") or "").strip().lower() for field in merged["fields"]}
    unmapped: list[str] = []
    for field in _docx_fields_for_schema(schema, pdf_path):
        field_id = str(field["id"])
        label_key = str(field["label"]).strip().lower()
        if field_id in used_ids or label_key in existing_labels:
            continue
        if not _field_has_preview_rect(field):
            unmapped.append(str(field.get("label") or field_id))
            continue
        used_ids.add(field_id)
        existing_labels.add(label_key)
        merged["fields"].append(field)
    if unmapped:
        preview = ", ".join(unmapped[:8])
        suffix = "" if len(unmapped) <= 8 else f", +{len(unmapped) - 8} field lain"
        raise HTTPException(
            status_code=422,
            detail=(
                "Schema form belum lengkap: field DOCX berikut belum punya koordinat preview: "
                f"{preview}{suffix}. Regenerate schema form dulu."
            ),
        )
    return merged


def _schema_with_optional_docx_fields(schema: dict[str, Any], pdf_path: Path) -> dict[str, Any]:
    try:
        return _schema_with_docx_fields(schema, pdf_path)
    except HTTPException as error:
        if error.status_code == 500:
            return deepcopy(schema)
        raise


def _docx_value_for_field(value: Any, field: dict[str, Any]) -> str:
    if str(field.get("type")) == "checkbox":
        return "X" if _coerce_checkbox(value) else ""
    if isinstance(value, bool):
        return "X" if value else ""
    return str(value or "").strip()


def _replace_paragraph_placeholders(
    paragraph: Any,
    field_iter: Any,
    values: dict[str, str | bool],
) -> None:
    runs = list(paragraph.runs)
    text = "".join(run.text for run in runs) if runs else paragraph.text
    if not FORM_PLACEHOLDER_PATTERN.search(text):
        return

    output: list[str] = []
    cursor = 0
    changed = False
    for match in FORM_PLACEHOLDER_PATTERN.finditer(text):
        output.append(text[cursor : match.start()])
        field = next(field_iter, None)
        if field is None:
            output.append(match.group(0))
        else:
            output.append(_docx_value_for_field(values.get(str(field["id"])), field))
            changed = True
        cursor = match.end()
    output.append(text[cursor:])
    if not changed:
        return

    rendered = "".join(output)
    if runs:
        runs[0].text = rendered
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(rendered)


def _iter_docx_paragraphs_for_fill(document: Any):
    for kind, block in _iter_docx_blocks(document):
        if kind == "paragraph":
            yield block
            continue

        seen_cells: set[int] = set()
        for row in block.rows:
            for cell in row.cells:
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                yield from cell.paragraphs


def _fill_docx_blank_cells(
    document: Any,
    fields: list[dict[str, Any]],
    values: dict[str, str | bool],
) -> None:
    for field in fields:
        target = field.get("docx_target")
        if not isinstance(target, dict):
            continue
        value = _docx_value_for_field(values.get(str(field["id"])), field)
        if not value:
            continue
        try:
            table = document.tables[int(target["table"])]
            cell = table.cell(int(target["row"]), int(target["column"]))
        except (IndexError, KeyError, TypeError, ValueError):
            continue
        if cell.paragraphs:
            cell.paragraphs[0].text = value
        else:
            cell.add_paragraph(value)


def fill_docx_schema_form(
    path: Path,
    values: dict[str, str | bool],
    signature_files: dict[str, dict[str, Any]] | None = None,
) -> bytes:
    if signature_files:
        raise HTTPException(
            status_code=422,
            detail="Upload tanda tangan belum didukung untuk download Word. Pilih PDF untuk memakai tanda tangan.",
        )

    schema = _schema_for_resolved_path(path)
    if schema is None:
        raise HTTPException(status_code=404, detail="Schema editor belum tersedia untuk form ini.")
    base_schema = schema
    merged_schema = _schema_with_docx_fields(base_schema, path)
    field_map = _schema_field_map(merged_schema)
    unknown_values = sorted(key for key in values if key not in field_map)
    if unknown_values:
        raise HTTPException(
            status_code=422,
            detail=f"Field tidak dikenal: {', '.join(unknown_values)}.",
        )

    return fill_docx_template_placeholders(
        word_template_docx_path(path),
        values,
        fields=_docx_fields_for_schema(base_schema, path),
    )


def fill_docx_template_placeholders(
    docx_path: Path,
    values: dict[str, str | bool],
    *,
    fields: list[dict[str, Any]] | None = None,
) -> bytes:
    fields = fields or extract_docx_placeholder_fields(docx_path)
    document = _load_word_document(docx_path)
    field_iter = iter([field for field in fields if field.get("docx_placeholder")])
    for paragraph in _iter_docx_paragraphs_for_fill(document):
        _replace_paragraph_placeholders(paragraph, field_iter, values)
    _fill_docx_blank_cells(document, fields, values)

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "filled-form.docx"
        document.save(str(output_path))
        content = output_path.read_bytes()
    if not content:
        raise HTTPException(status_code=500, detail="Form Word gagal dibuat.")
    return content


def _segment_label(field: dict, segments: list[dict]) -> str | None:
    """Cari label teks bersih untuk satu segmen placeholder.

    Urutan pencarian: isi bracket dulu, lalu teks terdekat
    di kiri pada baris yang sama, lalu teks tepat di atasnya. Jika tidak ada
    label yang layak, kembalikan None agar segmen itu dilewati.
    """
    inside = field["text"][1:-1].strip()
    if inside:
        return inside

    fx0, fy0, fx1, fy1 = field["bbox"]
    field_center_y = (fy0 + fy1) / 2

    best_left, best_left_x1 = None, float("-inf")
    for segment in segments:
        if segment is field or FORM_FIELD_CELL_PATTERN.match(segment["text"]):
            continue
        x0, y0, x1, y1 = segment["bbox"]
        if y0 <= field_center_y <= y1 and x1 <= fx0 + 2 and x1 > best_left_x1:
            candidate = _clean_label(segment["text"])
            if candidate:
                best_left, best_left_x1 = candidate, x1
    if best_left:
        return best_left

    best_above, best_above_y1 = None, float("-inf")
    for segment in segments:
        if segment is field or FORM_FIELD_CELL_PATTERN.match(segment["text"]):
            continue
        x0, y0, x1, y1 = segment["bbox"]
        if y1 <= fy0 + 2 and not (x1 < fx0 or x0 > fx1) and y1 > best_above_y1:
            candidate = _clean_label(segment["text"])
            if candidate:
                best_above, best_above_y1 = candidate, y1
    return best_above


def _cluster_rows(segments: list[dict]) -> list[dict]:
    """Kelompokkan segmen menjadi baris berdasarkan tumpang tindih vertikal."""
    rows: list[dict] = []
    for segment in sorted(segments, key=lambda s: (round(s["bbox"][1]), s["bbox"][0])):
        top, bottom = segment["bbox"][1], segment["bbox"][3]
        if rows and top < rows[-1]["bottom"] - 1:
            rows[-1]["items"].append(segment)
            rows[-1]["bottom"] = max(rows[-1]["bottom"], bottom)
        else:
            rows.append({"items": [segment], "bottom": bottom})
    return rows


def _field_segments(doc) -> list[dict]:
    """Ambil field placeholder dari blok info atas tiap halaman.

    Hanya blok baris placeholder pertama yang diambil. Bagian lanjutan seperti
    tabel log, teks bebas, atau blok tanda tangan dilewati agar form tetap singkat.
    """
    fields: list[dict] = []
    for page_number, page in enumerate(doc):
        segments = _page_segments(page, page_number)
        started = False
        for row in _cluster_rows(segments):
            row_fields: list[dict] = []
            for segment in sorted(row["items"], key=lambda s: s["bbox"][0]):
                if not FORM_FIELD_CELL_PATTERN.match(segment["text"]):
                    continue
                label = _segment_label(segment, segments)
                if label:
                    x0, y0, x1, y1 = segment["bbox"]
                    row_fields.append(
                        {
                            "key": f"{page_number}:{x0:.0f}:{y0:.0f}:{x1:.0f}:{y1:.0f}",
                            "label": label,
                            "page": page_number,
                            "bbox": segment["bbox"],
                            "size": segment["size"],
                            "row": row["items"],
                        }
                    )
            if row_fields:
                started = True
                fields.extend(row_fields)
            elif started:
                break  # end of the top info block for this page
    return fields


def _scan_form_fields(path: Path) -> list[dict[str, str]]:
    # Kembalikan daftar field {key, label} dari template form PDF.
    with fitz.open(path) as doc:
        return [{"key": field["key"], "label": field["label"]} for field in _field_segments(doc)]


def _unique_form_fields(path: Path) -> list[dict[str, str]]:
    """Hilangkan field duplikat berdasarkan label agar input tidak berulang."""
    seen: set[str] = set()
    unique: list[dict[str, str]] = []
    for field in _scan_form_fields(path):
        if field["label"] in seen:
            continue
        seen.add(field["label"])
        unique.append({"key": field["label"], "label": field["label"]})
    return unique


def _fit_font_size(page, field: dict, value: str) -> float:
    """Perkecil font agar nilai muat sampai batas kanan placeholder."""
    x0, _y0, x1, y1 = field["bbox"]
    field_center_y = (field["bbox"][1] + y1) / 2
    right_boundary = page.rect.width - 40
    for segment in field["row"]:
        sx0, sy0, _sx1, sy1 = segment["bbox"]
        if sy0 <= field_center_y <= sy1 and sx0 >= x1 - 1:
            right_boundary = min(right_boundary, sx0 - 2)
    available = max(right_boundary - x0, 20)

    size = field["size"]
    while size > _MIN_FONT_SIZE and _text_length(value, fontsize=size) > available:
        size -= 0.5
    return size


def _fill_form_placeholders(path: Path, values: dict[str, str]) -> bytes:
    """Isi setiap field yang labelnya punya nilai lalu kembalikan PDF hasilnya.

    Nilai dipetakan berdasarkan label, jadi satu input bisa mengisi semua field
    dengan label yang sama. Nilai kosong dibiarkan apa adanya.
    """
    doc = fitz.open(path)
    try:
        fields_by_label: dict[str, list[dict]] = {}
        for field in _field_segments(doc):
            fields_by_label.setdefault(field["label"], []).append(field)

        for label, raw_value in values.items():
            value = str(raw_value).strip()
            if not value:
                continue
            for field in fields_by_label.get(label, []):
                page = doc[field["page"]]
                x0, y0, x1, y1 = field["bbox"]
                size = _fit_font_size(page, field, value)
                # Hapus placeholder lama dari content PDF, bukan cuma ditutup putih.
                page.add_redact_annot(fitz.Rect(x0, y0, x1, y1), fill=(1, 1, 1))
                page.apply_redactions()
                page.insert_text(
                    (x0, y1 - 2),
                    value,
                    fontname=_FILL_FONT,
                    fontfile=_FILL_FONT_FILE,
                    fontsize=size,
                    color=(0, 0, 0),
                )

        return doc.tobytes(deflate=True)
    finally:
        doc.close()


def _resolve_form_path(path: str) -> Path:
    # Tentukan dan validasi bahwa path menunjuk ke template form PDF yang bisa diisi.
    resolved_path = _resolve_document_path(path)
    if not resolved_path.exists() or not resolved_path.is_file():
        raise HTTPException(status_code=404, detail="Form not found.")
    if (
        resolved_path.suffix.lower() != ".pdf"
        or _document_kind_for_path(resolved_path) != "form"
    ):
        raise HTTPException(status_code=400, detail="Dokumen ini bukan form yang bisa diisi.")
    return resolved_path


def _form_schema_dir() -> Path:
    return Path(__file__).resolve().parent.parent / "form_schemas"


def _relative_form_path(path: Path) -> str:
    return path.resolve().relative_to(_get_data_dir().resolve()).as_posix()


def _coerce_checkbox(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on"}
    if isinstance(value, (int, float)):
        return bool(value)
    return False


def _rect_from_schema(rect: dict[str, Any]) -> fitz.Rect:
    x = float(rect["x"])
    y = float(rect["y"])
    width = float(rect["width"])
    height = float(rect["height"])
    return fitz.Rect(x, y, x + width, y + height)


def _field_has_preview_rect(field: dict[str, Any]) -> bool:
    rect = field.get("rect")
    return (
        field.get("page") is not None
        and isinstance(rect, dict)
        and all(key in rect for key in ("x", "y", "width", "height"))
    )


def _public_field(field: dict[str, Any]) -> dict[str, Any]:
    rect = field.get("rect")
    if not _field_has_preview_rect(field):
        raise HTTPException(
            status_code=422,
            detail=f'Field "{field.get("label") or field.get("id")}" belum punya koordinat preview.',
        )
    payload = {
        "id": str(field["id"]),
        "label": str(field["label"]),
        "type": str(field["type"]),
        "page": int(field["page"]),
        "rect": {
            "x": float(rect["x"]),
            "y": float(rect["y"]),
            "width": float(rect["width"]),
            "height": float(rect["height"]),
        },
        "required": bool(field.get("required", False)),
        "section": str(field.get("section") or ""),
        "placeholder": str(field.get("placeholder") or "") or None,
        "font_size": float(field.get("font_size") or 10),
        "align": str(field.get("align") or "left"),
        "line_height": float(field.get("line_height") or 1.08),
        "clear": bool(field.get("clear", True)),
        "clear_padding": max(float(field.get("clear_padding") or 1.0), 0.0),
    }
    layout = field.get("layout")
    if isinstance(layout, dict):
        payload["layout"] = {str(key): str(value) for key, value in layout.items() if value is not None}
    return payload


def _validate_schema_payload(payload: dict[str, Any], schema_path: Path) -> dict[str, Any]:
    if not isinstance(payload, dict):
        raise HTTPException(status_code=500, detail=f"Schema {schema_path.name} tidak valid.")
    path = str(payload.get("path") or "").strip()
    title = str(payload.get("title") or "").strip()
    pages = payload.get("pages")
    fields = payload.get("fields")
    if not path or not title or not isinstance(pages, list) or not isinstance(fields, list):
        raise HTTPException(status_code=500, detail=f"Schema {schema_path.name} tidak lengkap.")

    normalized_pages: list[dict[str, Any]] = []
    for page in pages:
        if not isinstance(page, dict):
            raise HTTPException(status_code=500, detail=f"Schema {schema_path.name} tidak valid.")
        normalized_pages.append(
            {
                "number": int(page["number"]),
                "width": float(page["width"]),
                "height": float(page["height"]),
            }
        )
    page_numbers = {page["number"] for page in normalized_pages}

    normalized_fields: list[dict[str, Any]] = []
    seen_ids: set[str] = set()
    for field in fields:
        if not isinstance(field, dict):
            raise HTTPException(status_code=500, detail=f"Schema {schema_path.name} tidak valid.")
        field_id = str(field.get("id") or "").strip()
        field_type = str(field.get("type") or "").strip()
        page_number = int(field.get("page"))
        rect = field.get("rect")
        if (
            not field_id
            or field_id in seen_ids
            or field_type not in FORM_FIELD_TYPES
            or page_number not in page_numbers
            or not isinstance(rect, dict)
        ):
            raise HTTPException(status_code=500, detail=f"Schema {schema_path.name} tidak valid.")
        seen_ids.add(field_id)
        normalized_field = {
            "id": field_id,
            "label": str(field.get("label") or field_id),
            "type": field_type,
            "page": page_number,
            "rect": {
                "x": float(rect["x"]),
                "y": float(rect["y"]),
                "width": float(rect["width"]),
                "height": float(rect["height"]),
            },
            "required": bool(field.get("required", False)),
            "section": str(field.get("section") or ""),
            "placeholder": str(field.get("placeholder") or "") or None,
            "font_size": float(field.get("font_size") or 10),
            "align": str(field.get("align") or "left"),
            "line_height": float(field.get("line_height") or 1.08),
            "clear": bool(field.get("clear", True)),
            "clear_padding": max(float(field.get("clear_padding") or 1.0), 0.0),
        }
        layout = field.get("layout")
        if isinstance(layout, dict):
            normalized_field["layout"] = {
                str(key): str(value)
                for key, value in layout.items()
                if key in {"kind", "group_id", "group_label", "row_label", "column_label", "choice_group"}
                and value is not None
            }
        normalized_fields.append(normalized_field)

    normalized_payload = {
        "path": path,
        "title": title,
        "pages": normalized_pages,
        "fields": normalized_fields,
    }
    generator = payload.get("generator")
    if isinstance(generator, dict):
        raw_warnings = generator.get("warnings", [])
        warnings = raw_warnings if isinstance(raw_warnings, list) else []
        normalized_payload["generator"] = {
            "source": str(generator.get("source") or ""),
            "quality_score": float(generator.get("quality_score") or 0),
            "warnings": [str(warning) for warning in warnings],
        }
    return normalized_payload


def _load_form_schema_payloads() -> list[dict[str, Any]]:
    schema_dir = _form_schema_dir()
    if not schema_dir.exists():
        return []

    payloads: list[dict[str, Any]] = []
    for schema_path in sorted(schema_dir.glob("*.json")):
        payload = json.loads(schema_path.read_text(encoding="utf-8"))
        payloads.append(_validate_schema_payload(payload, schema_path))
    return payloads


def _schema_for_resolved_path(path: Path) -> dict[str, Any] | None:
    relative_path = _relative_form_path(path)
    for payload in _load_form_schema_payloads():
        if payload["path"] == relative_path:
            return payload
    return None


def _schema_field_map(schema: dict[str, Any]) -> dict[str, dict[str, Any]]:
    return {str(field["id"]): field for field in schema["fields"]}


def has_schema_form(path: Path) -> bool:
    return _schema_for_resolved_path(path) is not None


def get_form_schema(path: str) -> dict[str, Any]:
    resolved_path = _resolve_form_path(path)
    schema = _schema_for_resolved_path(resolved_path)
    if schema is None:
        raise HTTPException(status_code=404, detail="Schema editor belum tersedia untuk form ini.")
    schema = _schema_with_docx_fields(schema, resolved_path)
    return {
        "path": schema["path"],
        "title": schema["title"],
        "pages": schema["pages"],
        "fields": [_public_field(field) for field in schema["fields"]],
    }


def _shape_text_fits(
    page: fitz.Page,
    rect: fitz.Rect,
    value: str,
    *,
    font_size: float,
    align: str,
    line_height: float,
) -> fitz.Shape | None:
    shape = page.new_shape()
    spare = shape.insert_textbox(
        rect,
        value,
        fontname=_FILL_FONT,
        fontfile=_FILL_FONT_FILE,
        fontsize=font_size,
        color=(0, 0, 0),
        align=_ALIGNMENTS.get(align, 0),
        lineheight=line_height,
    )
    return shape if spare >= 0 else None


def _clear_rect_for_field(field: dict[str, Any]) -> fitz.Rect:
    rect = _rect_from_schema(field["rect"])
    padding = max(float(field.get("clear_padding") or 0.0), 0.0)
    if padding <= 0:
        return rect
    return fitz.Rect(
        rect.x0 - padding,
        rect.y0 - padding,
        rect.x1 + padding,
        rect.y1 + padding,
    )


def _clear_field_content(page: fitz.Page, field: dict[str, Any]) -> None:
    if not field.get("clear", True):
        return
    page.add_redact_annot(_clear_rect_for_field(field), fill=(1, 1, 1))
    page.apply_redactions()


def _write_text_field(page: fitz.Page, field: dict[str, Any], value: str) -> None:
    rect = _rect_from_schema(field["rect"])
    font_size = float(field.get("font_size") or 10)
    line_height = float(field.get("line_height") or 1.08)
    align = str(field.get("align") or "left")

    chosen_size: float | None = None
    current_size = font_size
    while current_size >= _MIN_FONT_SIZE:
        if _shape_text_fits(
            page,
            rect,
            value,
            font_size=current_size,
            align=align,
            line_height=line_height,
        ) is not None:
            chosen_size = current_size
            break
        current_size -= 0.5

    _clear_field_content(page, field)

    chosen_shape = page.new_shape()
    chosen_shape.insert_textbox(
        rect,
        value,
        fontname=_FILL_FONT,
        fontfile=_FILL_FONT_FILE,
        fontsize=chosen_size or _MIN_FONT_SIZE,
        color=(0, 0, 0),
        align=_ALIGNMENTS.get(align, 0),
        lineheight=line_height,
    )
    chosen_shape.commit(overlay=True)


def _write_checkbox(page: fitz.Page, field: dict[str, Any]) -> None:
    rect = _rect_from_schema(field["rect"])
    _clear_field_content(page, field)

    font_size = max(min(rect.height * 0.95, 16), _MIN_FONT_SIZE)
    text_width = _text_length("X", fontsize=font_size)
    page.insert_text(
        (
            rect.x0 + max((rect.width - text_width) / 2, 0),
            rect.y1 - max((rect.height - font_size) / 2.8, 1),
        ),
        "X",
        fontname=_FILL_FONT,
        fontfile=_FILL_FONT_FILE,
        fontsize=font_size,
        color=(0, 0, 0),
    )


def _write_signature_image(page: fitz.Page, field: dict[str, Any], image_content: bytes) -> None:
    rect = _rect_from_schema(field["rect"])
    _clear_field_content(page, field)
    page.insert_image(rect, stream=image_content, keep_proportion=True, overlay=True)


def _validate_schema_values(
    schema: dict[str, Any],
    values: dict[str, str | bool],
    signature_files: dict[str, dict[str, Any]],
) -> None:
    field_map = _schema_field_map(schema)
    unknown_values = sorted(key for key in values if key not in field_map)
    if unknown_values:
        raise HTTPException(
            status_code=422,
            detail=f"Field tidak dikenal: {', '.join(unknown_values)}.",
        )

    invalid_signature_keys = sorted(
        key
        for key in signature_files
        if key not in field_map or field_map[key]["type"] != "signature_image"
    )
    if invalid_signature_keys:
        raise HTTPException(
            status_code=422,
            detail=f"Upload signature tidak valid untuk field: {', '.join(invalid_signature_keys)}.",
        )


def fill_schema_form(
    path: Path,
    values: dict[str, str | bool],
    signature_files: dict[str, dict[str, Any]],
) -> bytes:
    schema = _schema_for_resolved_path(path)
    if schema is None:
        raise HTTPException(status_code=404, detail="Schema editor belum tersedia untuk form ini.")

    for field_id, file_payload in signature_files.items():
        content_type = str(file_payload.get("content_type") or "").lower()
        if content_type not in IMAGE_CONTENT_TYPES:
            raise HTTPException(
                status_code=422,
                detail=f'Field "{field_id}" harus berupa PNG atau JPEG.',
            )
        if not isinstance(file_payload.get("content"), bytes) or not file_payload["content"]:
            raise HTTPException(
                status_code=422,
                detail=f'File untuk field "{field_id}" kosong.',
            )

    _validate_schema_values(schema, values, signature_files)

    doc = fitz.open(path)
    try:
        for field in schema["fields"]:
            if field.get("page") is None or not isinstance(field.get("rect"), dict):
                continue
            page = doc[int(field["page"])]
            field_id = str(field["id"])
            field_type = str(field["type"])
            if field_type == "signature_image":
                upload = signature_files.get(field_id)
                if upload is None:
                    continue
                _write_signature_image(page, field, upload["content"])
                continue
            if field_type == "checkbox":
                if _coerce_checkbox(values.get(field_id)):
                    _write_checkbox(page, field)
                continue
            value = str(values.get(field_id) or "").strip()
            if not value:
                continue
            _write_text_field(page, field, value)
        return doc.tobytes(deflate=True)
    finally:
        doc.close()
